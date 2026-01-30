"""
Resume Generator - Creates professional Word documents (.docx) from resume data

Usage:
    python resume-generator.py <json_file>              # Load from JSON file
    python resume-generator.py                          # Create sample resumes
    python resume-generator.py --interactive            # Interactive mode
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import json
from datetime import datetime
import sys
import os

# ============================================================================
# STYLING CONFIGURATION - Tweak these values to change document appearance
# ============================================================================

# MARGIN SETTINGS (in inches)
MARGIN_TOP = 0.5        # Top margin - TWEAK HERE for more/less top space
MARGIN_BOTTOM = 0.5     # Bottom margin
MARGIN_LEFT = 0.5       # Left margin
MARGIN_RIGHT = 0.5      # Right margin

# PAPER SIZE (A4)
PAPER_WIDTH = 8.27      # A4 width in inches
PAPER_HEIGHT = 11.69    # A4 height in inches

# FONT SETTINGS
DEFAULT_FONT = 'Times New Roman'  # Change to other fonts like 'Calibri', 'Arial'
DEFAULT_FONT_SIZE = 11            # Body text font size in points - TWEAK HERE

# HEADER STYLING
HEADER_NAME_SIZE = 22             # Name font size in points - TWEAK HERE
HEADER_NAME_BOLD = True           # Make name bold
HEADER_CONTACT_SIZE = 11          # Contact info font size - TWEAK HERE
HEADER_CONTACT_SPACE_BEFORE = 3   # Space above contact line - TWEAK HERE
HEADER_CONTACT_SPACE_AFTER = 12   # Space below contact line - TWEAK HERE

# SECTION TITLE STYLING
SECTION_TITLE_SIZE = 14           # Section title font size - TWEAK HERE
SECTION_TITLE_BOLD = True         # Make section titles bold
SECTION_TITLE_SPACE_BEFORE = 6    # Space before section title - TWEAK HERE (you mentioned 6pts)
SECTION_TITLE_SPACE_AFTER = 6     # Space after section title - TWEAK HERE
SECTION_TITLE_BORDER_COLOR = '000000'  # Border color (000000 = black) - TWEAK HERE

# CONTENT SPACING
CONTENT_SPACE_AFTER = 6           # Space after job/project titles - TWEAK HERE
BULLET_ITEM_SPACE = 0             # Space after each bullet point - TWEAK HERE

# TABLE STYLING
SHOW_TABLE_BORDERS = True        # Show/hide table borders in education section - TWEAK HERE

# ALIGNMENT
DEFAULT_ALIGNMENT = WD_ALIGN_PARAGRAPH.JUSTIFY  # JUSTIFY, LEFT, CENTER, RIGHT
HEADER_ALIGNMENT = WD_ALIGN_PARAGRAPH.CENTER    # Center header

# ============================================================================


class ResumeGenerator:
    """Generate professional resumes in Word (.docx) format"""
    
    def __init__(self):
        self.doc = Document()
        self.set_document_margins()
        self.set_paper_size()
        self.set_default_styles()
    
    def set_document_margins(self):
        """Set document margins (in inches)"""
        # TWEAK: Modify MARGIN_* values at the top to change margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(MARGIN_TOP)
            section.bottom_margin = Inches(MARGIN_BOTTOM)
            section.left_margin = Inches(MARGIN_LEFT)
            section.right_margin = Inches(MARGIN_RIGHT)
    
    def set_paper_size(self):
        """Set paper size to A4"""
        # TWEAK: Change PAPER_WIDTH and PAPER_HEIGHT for different sizes
        from docx.shared import Inches
        sections = self.doc.sections
        for section in sections:
            section.page_height = Inches(PAPER_HEIGHT)
            section.page_width = Inches(PAPER_WIDTH)
    
    def set_default_styles(self):
        """Set up default font styles for the entire document"""
        # TWEAK: Change DEFAULT_FONT and DEFAULT_FONT_SIZE at the top
        style = self.doc.styles['Normal']
        style.font.name = DEFAULT_FONT
        style.font.size = Pt(DEFAULT_FONT_SIZE)
    
    def add_section_title(self, title):
        """Add a section title with bottom border and specified font size"""
        # Section title styling
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        # TWEAK: Modify SECTION_TITLE_SIZE to make titles bigger/smaller
        title_run.font.size = Pt(SECTION_TITLE_SIZE)
        title_run.font.bold = SECTION_TITLE_BOLD
        title_run.font.name = DEFAULT_FONT
        title_para.alignment = DEFAULT_ALIGNMENT
        # TWEAK: Modify SECTION_TITLE_SPACE_* values to adjust spacing
        title_para.paragraph_format.space_before = Pt(SECTION_TITLE_SPACE_BEFORE)
        title_para.paragraph_format.space_after = Pt(SECTION_TITLE_SPACE_AFTER)
        
        # Add bottom border to section title (single black line)
        pPr = title_para._element.get_or_add_pPr()
        pBdr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
        if pBdr is not None:
            pPr.remove(pBdr)
        
        # TWEAK: Border styling - change w:sz="12" (thickness), w:color="000000" (black)
        from docx.oxml import parse_xml
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/></w:pBdr>')
        pPr.append(pBdr)
        
        return title_para
    
    def add_header(self, name, email, phone, location):
        """Add header with name and contact information"""
        # ===== NAME SECTION =====
        name_para = self.doc.add_paragraph()
        name_run = name_para.add_run(name)
        # TWEAK: Modify HEADER_NAME_SIZE to change name font size
        name_run.font.size = Pt(HEADER_NAME_SIZE)
        name_run.font.bold = HEADER_NAME_BOLD
        name_run.font.name = DEFAULT_FONT
        name_para.alignment = HEADER_ALIGNMENT
        
        # ===== CONTACT INFO SECTION =====
        contact_para = self.doc.add_paragraph()
        contact_para.alignment = HEADER_ALIGNMENT
        contact_text = f"Phone: {phone} | Email: {email} | {location}"
        contact_run = contact_para.add_run(contact_text)
        contact_run.font.name = DEFAULT_FONT
        # TWEAK: Modify HEADER_CONTACT_SIZE to change contact info font size
        contact_run.font.size = Pt(HEADER_CONTACT_SIZE)
        # TWEAK: Modify HEADER_CONTACT_SPACE_* to adjust spacing around contact
        contact_para.paragraph_format.space_before = Pt(HEADER_CONTACT_SPACE_BEFORE)
        contact_para.paragraph_format.space_after = Pt(HEADER_CONTACT_SPACE_AFTER)
        
        # Add spacing paragraph after header
        self.doc.add_paragraph()
    
    def add_section(self, title, content_list, use_bullets=False):
        """Add a resume section (education, experience, projects, etc.)"""
        # Add section title with border
        self.add_section_title(title)
        
        # ===== CONTENT ITEMS =====
        for item in content_list:
            if isinstance(item, dict):
                # ===== STRUCTURED ENTRIES (job, project, education) =====
                position_para = self.doc.add_paragraph()
                position_para.alignment = DEFAULT_ALIGNMENT
                
                # Job title or project title (BOLD)
                if 'title' in item:
                    title_run = position_para.add_run(item['title'])
                    title_run.bold = True
                    title_run.font.name = DEFAULT_FONT
                    # TWEAK: Change Pt(11) to make job titles bigger/smaller
                    title_run.font.size = Pt(11)
                
                # Organization and dates (normal)
                if 'organization' in item or 'dates' in item:
                    separator = " | " if 'title' in item else ""
                    org_dates = f"{item.get('organization', '')}"
                    if 'dates' in item:
                        org_dates += f" ({item['dates']})"
                    
                    org_run = position_para.add_run(separator + org_dates)
                    org_run.font.name = DEFAULT_FONT
                    # TWEAK: Change Pt(11) to adjust organization/dates font size
                    org_run.font.size = Pt(11)
                
                # TWEAK: Modify CONTENT_SPACE_AFTER to change space after job titles
                position_para.paragraph_format.space_after = Pt(CONTENT_SPACE_AFTER)
                
                # ===== BULLET POINT DESCRIPTIONS =====
                if 'description' in item:
                    if isinstance(item['description'], list):
                        for desc in item['description']:
                            desc_para = self.doc.add_paragraph(desc, style='List Bullet')
                            desc_para.alignment = DEFAULT_ALIGNMENT
                            for run in desc_para.runs:
                                run.font.name = DEFAULT_FONT
                                # TWEAK: Change Pt(11) to adjust bullet point font size
                                run.font.size = Pt(11)
                            # TWEAK: Modify BULLET_ITEM_SPACE to change space between bullets
                            desc_para.paragraph_format.space_after = Pt(BULLET_ITEM_SPACE)
                    else:
                        desc_para = self.doc.add_paragraph(item['description'], style='List Bullet')
                        desc_para.alignment = DEFAULT_ALIGNMENT
                        for run in desc_para.runs:
                            run.font.name = DEFAULT_FONT
                            run.font.size = Pt(11)
            else:
                # ===== SIMPLE TEXT ENTRIES (professional summary, etc.) =====
                if use_bullets:
                    # Render as bullet points
                    text_para = self.doc.add_paragraph(item, style='List Bullet')
                else:
                    # Render as regular paragraph
                    text_para = self.doc.add_paragraph(item)
                text_para.alignment = DEFAULT_ALIGNMENT
                for run in text_para.runs:
                    run.font.name = DEFAULT_FONT
                    # TWEAK: Change Pt(11) to adjust text font size
                    run.font.size = Pt(11)
    
    def add_skills_section(self, skills_data):
        """Add skills section with bold categories and normal text content"""
        # Add section title with border
        self.add_section_title("TECHNICAL SKILLS")
        
        # ===== SKILLS ITEMS =====
        for skill_item in skills_data:
            skill_para = self.doc.add_paragraph()
            # TWEAK: Change DEFAULT_ALIGNMENT to LEFT or CENTER for different alignment
            skill_para.alignment = DEFAULT_ALIGNMENT
            
            # Split by colon to separate category from content
            # Format: "Category: content here"
            if ':' in skill_item:
                category, content = skill_item.split(':', 1)
                
                # ===== BOLD CATEGORY PART =====
                category_run = skill_para.add_run(category.strip() + ':')
                category_run.bold = True  # TWEAK: Set to False to make category normal
                category_run.font.name = DEFAULT_FONT
                # TWEAK: Change Pt(11) to adjust category font size
                category_run.font.size = Pt(11)
                
                # ===== NORMAL CONTENT PART =====
                content_run = skill_para.add_run(' ' + content.strip())
                content_run.bold = False  # TWEAK: Set to True to make content bold
                content_run.font.name = DEFAULT_FONT
                # TWEAK: Change Pt(11) to adjust content font size
                content_run.font.size = Pt(11)
            else:
                # If no colon, treat entire line as bold
                skill_run = skill_para.add_run(skill_item)
                skill_run.bold = True
                skill_run.font.name = DEFAULT_FONT
                skill_run.font.size = Pt(11)
            
            # TWEAK: Modify BULLET_ITEM_SPACE to change space between skill lines
            skill_para.paragraph_format.space_after = Pt(BULLET_ITEM_SPACE)
    
    def add_education_table(self, education_data):
        """Add education section as a table with columns for organization/degree and dates/location"""
        # Add section title with border
        self.add_section_title("EDUCATION")
        
        # ===== CREATE TABLE FOR EDUCATION =====
        # 2 columns: Left column for school/degree, Right column for dates/location
        for edu in education_data:
            # Create 2x2 table (2 rows, 2 columns)
            table = self.doc.add_table(rows=2, cols=2)
            # TWEAK: Change table.style to 'Table Grid', 'Light Grid', etc.
            table.style = 'Table Grid'
            
            # ===== FIRST ROW: Organization (left) | Dates (right) =====
            # Left cell - School/University name
            cell_org = table.rows[0].cells[0]
            org_para = cell_org.paragraphs[0]
            org_run = org_para.add_run(edu.get('organization', ''))
            org_run.font.name = DEFAULT_FONT
            # TWEAK: Change Pt(11) to adjust organization font size
            org_run.font.size = Pt(11)
            org_run.font.bold = False
            org_para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left align org
            
            # Right cell - Dates
            cell_dates = table.rows[0].cells[1]
            dates_para = cell_dates.paragraphs[0]
            dates_run = dates_para.add_run(edu.get('dates', ''))
            dates_run.font.name = DEFAULT_FONT
            # TWEAK: Change Pt(11) to adjust dates font size
            dates_run.font.size = Pt(11)
            dates_run.font.bold = False
            dates_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Right align dates
            
            # ===== SECOND ROW: Degree (left) | Location (right) =====
            # Left cell - Degree/Program name
            cell_degree = table.rows[1].cells[0]
            degree_para = cell_degree.paragraphs[0]
            degree_run = degree_para.add_run(edu.get('title', ''))
            degree_run.font.name = DEFAULT_FONT
            # TWEAK: Change Pt(11) to adjust degree font size
            degree_run.font.size = Pt(11)
            degree_run.font.bold = False
            degree_para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left align degree
            
            # Right cell - Location
            cell_location = table.rows[1].cells[1]
            location_para = cell_location.paragraphs[0]
            location_run = location_para.add_run(edu.get('location', ''))
            location_run.font.name = DEFAULT_FONT
            # TWEAK: Change Pt(11) to adjust location font size
            location_run.font.size = Pt(11)
            location_run.font.bold = False
            location_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Right align location
            
            # ===== REMOVE TABLE BORDERS =====
            # TWEAK: Set SHOW_TABLE_BORDERS to True at the top to show visible table borders
            if not SHOW_TABLE_BORDERS:
                # Remove all borders for invisible table effect
                from docx.oxml import parse_xml
                for row in table.rows:
                    for cell in row.cells:
                        tcPr = cell._element.get_or_add_tcPr()
                        tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tcBorders>')
                        tcPr.append(tcBorders)
            
            # TWEAK: Modify this value to change space after education entry
            spacing_para = self.doc.add_paragraph()
            spacing_para.paragraph_format.space_after = Pt(3)
    
    def save(self, filename='resume.docx'):
        """Save the resume to a Word document"""
        output_path = Path(filename)
        self.doc.save(output_path)
        return str(output_path.absolute())


def create_sample_resume():
    """Create a sample resume to demonstrate the generator"""
    generator = ResumeGenerator()
    
    # Header
    generator.add_header(
        name="John Doe",
        email="john.doe@email.com",
        phone="(555) 123-4567",
        location="New York, NY"
    )
    
    # Professional Summary
    generator.add_section("PROFESSIONAL SUMMARY", [
        "Results-driven software engineer with 5+ years of experience developing scalable applications. "
        "Expertise in Python, web development, and cloud technologies."
    ])
    
    # Experience
    generator.add_section("PROFESSIONAL EXPERIENCE", [
        {
            "title": "Senior Software Engineer",
            "organization": "Tech Company Inc.",
            "dates": "Jan 2022 - Present",
            "description": [
                "Led development of microservices architecture serving 1M+ users",
                "Improved application performance by 40% through optimization",
                "Mentored junior developers and conducted code reviews"
            ]
        },
        {
            "title": "Software Developer",
            "organization": "Digital Solutions Ltd.",
            "dates": "Jun 2019 - Dec 2021",
            "description": [
                "Developed full-stack web applications using Python and React",
                "Implemented REST APIs and database optimization",
                "Collaborated with product team on feature specifications"
            ]
        }
    ])
    
    # Projects
    generator.add_section("PROJECTS", [
        {
            "title": "Cloud Infrastructure Automation",
            "description": [
                "Built automated deployment pipeline using Terraform and Python",
                "Reduced infrastructure setup time by 70%"
            ]
        },
        {
            "title": "Real-time Data Pipeline",
            "description": [
                "Designed and implemented data processing system handling 1M+ events/day",
                "Implemented monitoring and alerting for production reliability"
            ]
        }
    ])
    
    # Skills
    generator.add_skills_section([
        "Languages: Python, JavaScript, SQL, Java",
        "Frameworks: Django, Flask, React, FastAPI",
        "Tools: Git, Docker, AWS, PostgreSQL",
        "Methodologies: Agile, RESTful API Design"
    ])
    
    # Education
    generator.add_education_table([
        {
            "title": "Bachelor of Science in Computer Science",
            "organization": "State University",
            "dates": "2019"
        }
    ])
    
    # Save
    output_file = generator.save('sample_resume.docx')
    print(f"✓ Sample resume created: {output_file}")
    return generator
def create_custom_resume(resume_data):
    """Create a resume from custom data
    
    Args:
        resume_data (dict): Dictionary containing resume information with keys:
            - personal: {name, email, phone, location}
            - summary: str
            - experience: list of dicts
            - education: list of dicts
            - skills: list of str
            - projects: list of dicts (optional)
    """
    generator = ResumeGenerator()
    
    # Header
    personal = resume_data.get('personal', {})
    generator.add_header(
        name=personal.get('name', 'Your Name'),
        email=personal.get('email', 'email@example.com'),
        phone=personal.get('phone', '(555) 000-0000'),
        location=personal.get('location', 'City, State')
    )
    
    # Professional Summary
    if resume_data.get('summary'):
        summary_list = resume_data['summary'] if isinstance(resume_data['summary'], list) else [resume_data['summary']]
        generator.add_section("PROFESSIONAL SUMMARY", summary_list, use_bullets=True)
    
    # Experience
    if resume_data.get('experience'):
        generator.add_section("PROFESSIONAL EXPERIENCE", resume_data['experience'])
    
    # Projects
    if resume_data.get('projects'):
        generator.add_section("PROJECTS", resume_data['projects'])
    
    # Skills
    if resume_data.get('skills'):
        generator.add_skills_section(resume_data['skills'])
    
    # Education (using table format)
    if resume_data.get('education'):
        generator.add_education_table(resume_data['education'])
    
    return generator


def main():
    """Main function to handle CLI arguments and different modes"""
    
    # Check for command-line arguments
    if len(sys.argv) > 1:
        arg = sys.argv[1]
        
        # Interactive mode
        if arg == '--interactive':
            interactive_mode()
            return
        
        # Load from JSON file
        if arg.endswith('.json'):
            load_from_json(arg)
            return
        
        # Help or unknown argument
        print("Usage:")
        print("  python resume-generator.py <file.json>    Load resume from JSON file")
        print("  python resume-generator.py --interactive   Interactive mode")
        print("  python resume-generator.py                 Create sample resumes")
        return
    
    # Default: Create sample resumes
    print("Resume Generator - Word Document Creator")
    print("=" * 50)
    
    create_sample_resume()
    
    # Example: Create custom resume from JSON data
    custom_data = {
        "personal": {
            "name": "Jane Smith",
            "email": "jane.smith@email.com",
            "phone": "(555) 987-6543",
            "location": "Los Angeles, CA"
        },
        "summary": "Creative and analytical data scientist with 3+ years of experience in machine learning and data analysis.",
        "experience": [
            {
                "title": "Data Scientist",
                "organization": "Analytics Pro",
                "dates": "Mar 2022 - Present",
                "description": [
                    "Developed ML models for customer churn prediction",
                    "Automated data pipeline reducing processing time by 60%",
                    "Created interactive dashboards using Python and Tableau"
                ]
            },
            {
                "title": "Junior Data Analyst",
                "organization": "Business Insights",
                "dates": "Jan 2021 - Feb 2022",
                "description": [
                    "Analyzed large datasets to identify business trends",
                    "Created SQL queries for reporting and analysis"
                ]
            }
        ],
        "education": [
            {
                "title": "Master of Science in Data Science",
                "organization": "Tech University",
                "dates": "2020"
            },
            {
                "title": "Bachelor of Science in Statistics",
                "organization": "State College",
                "dates": "2019"
            }
        ],
        "skills": [
            "Python, R, SQL, Java",
            "Machine Learning, Statistical Analysis, Data Visualization",
            "TensorFlow, Scikit-learn, Pandas, NumPy",
            "Tableau, Power BI, Jupyter Notebooks"
        ]
    }
    
    generator = create_custom_resume(custom_data)
    output_file = generator.save('custom_resume.docx')
    print(f"✓ Custom resume created: {output_file}")


def load_from_json(json_file):
    """Load resume data from JSON file and generate Word document"""
    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        generator = create_custom_resume(data)
        
        # Create output folder if it doesn't exist
        output_dir = Path('output')
        output_dir.mkdir(exist_ok=True)
        
        # Format filename: FirstName_lastname_Resume_CompanyName.docx
        personal = data.get('personal', {})
        name_parts = personal.get('name', 'Resume').split()
        first_name = name_parts[0] if len(name_parts) > 0 else 'Resume'
        last_name = name_parts[-1].lower() if len(name_parts) > 1 else ''
        company_name = personal.get('company_name', 'Resume')
        
        if last_name and last_name != first_name.lower():
            output_name = f"{first_name}_{last_name}_Resume_{company_name}.docx"
        else:
            output_name = f"{first_name}_Resume_{company_name}.docx"
        
        output_path = output_dir / output_name
        
        # Check if file exists and create new version with counter if it does
        if output_path.exists():
            name_parts = output_name.rsplit('.', 1)
            base_name = name_parts[0]
            extension = '.' + name_parts[1] if len(name_parts) > 1 else ''
            counter = 1
            while output_path.exists():
                new_name = f"{base_name}_{counter}{extension}"
                output_path = output_dir / new_name
                counter += 1
        
        output_file = generator.save(str(output_path))
        print(f"✓ Resume created: {output_file}")
        
    except FileNotFoundError:
        print(f"Error: JSON file '{json_file}' not found.")
        sys.exit(1)
    except json.JSONDecodeError:
        print(f"Error: Invalid JSON format in '{json_file}'.")
        sys.exit(1)
    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)


def interactive_mode():
    """Interactive mode to build resume step by step"""
    print("Resume Generator - Interactive Mode")
    print("=" * 50)
    print()
    
    data = {}
    
    # Personal Information
    print("PERSONAL INFORMATION")
    print("-" * 50)
    data['personal'] = {
        'name': input("Full Name: "),
        'email': input("Email: "),
        'phone': input("Phone Number: "),
        'location': input("Location (City, State): ")
    }
    
    # Professional Summary
    print("\nPROFESSIONAL SUMMARY")
    print("-" * 50)
    data['summary'] = input("Brief professional summary (press Enter to skip): ").strip() or None
    
    # Experience
    print("\nEXPERIENCE")
    print("-" * 50)
    data['experience'] = []
    while True:
        if input("Add work experience? (y/n): ").lower() == 'y':
            exp = {
                'title': input("  Job Title: "),
                'organization': input("  Company/Organization: "),
                'dates': input("  Dates (e.g., Jan 2022 - Present): "),
                'description': []
            }
            print("  Enter job responsibilities (one per line, empty line to finish):")
            while True:
                desc = input("    - ").strip()
                if not desc:
                    break
                exp['description'].append(desc)
            data['experience'].append(exp)
        else:
            break
    
    # Education
    print("\nEDUCATION")
    print("-" * 50)
    data['education'] = []
    while True:
        if input("Add education? (y/n): ").lower() == 'y':
            edu = {
                'title': input("  Degree/Program: "),
                'organization': input("  School/University: "),
                'dates': input("  Graduation Year: ")
            }
            data['education'].append(edu)
        else:
            break
    
    # Skills
    print("\nSKILLS")
    print("-" * 50)
    data['skills'] = []
    print("Enter skills (one per line, empty line to finish):")
    while True:
        skill = input("  - ").strip()
        if not skill:
            break
        data['skills'].append(skill)
    
    # Save to JSON file
    print()
    json_filename = input("Save data to JSON file (default: resume_data.json): ").strip() or 'resume_data.json'
    with open(json_filename, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2)
    print(f"✓ Resume data saved to: {json_filename}")
    
    # Generate Word document
    generator = create_custom_resume(data)
    docx_filename = json_filename.replace('.json', '.docx')
    output_file = generator.save(docx_filename)
    print(f"✓ Resume generated: {output_file}")


if __name__ == "__main__":
    main()
